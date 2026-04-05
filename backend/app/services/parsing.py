"""Document parsing service for PDF, DOCX, XLSX formats."""
from typing import Dict, Any, Optional
import os
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook


class DocumentParser:
    """Parser for extracting text from various document formats."""

    SUPPORTED_FORMATS = ['pdf', 'docx', 'xlsx']

    @staticmethod
    def get_format(filename: str) -> Optional[str]:
        """Extract file format from filename."""
        if not filename or '.' not in filename:
            return None
        ext = filename.lower().split('.')[-1]
        return ext if ext in DocumentParser.SUPPORTED_FORMATS else None

    @staticmethod
    def parse(file_path: str, format_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse document and extract text + metadata.

        Args:
            file_path: Path to the document file
            format_type: Optional format override. If None, inferred from file_path

        Returns:
            Dict with 'text' and 'metadata' keys
        """
        # Infer format from file path if not provided
        if format_type is None:
            format_type = DocumentParser.get_format(file_path)

        parser_map = {
            'pdf': DocumentParser._parse_pdf,
            'docx': DocumentParser._parse_docx,
            'xlsx': DocumentParser._parse_xlsx
        }

        if format_type not in parser_map:
            raise ValueError(f"Unsupported format: {format_type}")

        return parser_map[format_type](file_path)

    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF."""
        reader = PdfReader(file_path)
        text_content = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                text_content.append(text)

        # Extract metadata
        metadata = {
            'pages': len(reader.pages),
            'page_count': len(reader.pages),
        }

        if reader.metadata:
            if reader.metadata.author:
                metadata['author'] = reader.metadata.author
            if reader.metadata.title:
                metadata['title'] = reader.metadata.title
            if reader.metadata.creator:
                metadata['creator'] = reader.metadata.creator

        return {
            'text': '\n\n'.join(text_content),
            'metadata': metadata
        }

    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from DOCX."""
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells if cell.text)
                if row_text.strip():
                    tables_text.append(row_text)

        all_text = paragraphs + tables_text
        combined_text = '\n\n'.join(all_text)

        metadata = {
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables),
            'words': len(combined_text.split()),
        }

        # Core properties
        if doc.core_properties:
            if doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            if doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            if doc.core_properties.created:
                metadata['created'] = str(doc.core_properties.created)

        return {
            'text': combined_text,
            'metadata': metadata
        }

    @staticmethod
    def _parse_xlsx(file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from XLSX."""
        wb = load_workbook(file_path, read_only=True, data_only=True)
        all_rows = []
        sheet_data = {}

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_rows = []

            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    sheet_rows.append(row_text)

            if sheet_rows:
                sheet_data[sheet_name] = len(sheet_rows)
                all_rows.append(f"=== Sheet: {sheet_name} ===")
                all_rows.extend(sheet_rows)

        combined_text = '\n'.join(all_rows)

        metadata = {
            'sheets': len(wb.sheetnames),
            'sheet_names': wb.sheetnames,
            'sheet_rows': sheet_data,
            'total_rows': len(all_rows),
        }

        return {
            'text': combined_text,
            'metadata': metadata
        }
