"""Generate sample test files for RAG pipeline manual testing.

Creates 3 files in docs/testing/fixtures/:
  - company-policy.pdf
  - financial-report-q4-2025.docx
  - product-catalog.xlsx

Usage:
    python docs/testing/generate-test-files.py
"""

import os
from pathlib import Path

# Output directory
FIXTURES_DIR = Path(__file__).parent / "fixtures"
FIXTURES_DIR.mkdir(exist_ok=True)


def generate_company_policy_pdf():
    """Generate a company policy PDF with remote work, leave, expenses, etc."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    path = FIXTURES_DIR / "company-policy.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    body = styles["Normal"]
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]

    elements = []

    # Title
    elements.append(Paragraph("Acme Corporation — Employee Handbook 2025", h1))
    elements.append(Paragraph("Effective Date: January 1, 2025 | Version 3.2", body))
    elements.append(Spacer(1, 20))

    # Section 1: Remote Work Policy
    elements.append(Paragraph("1. Remote Work Policy", h2))
    elements.append(Paragraph(
        "Acme Corporation allows eligible employees to work remotely up to <b>15 days per month</b>. "
        "Remote work must be approved by the employee's direct manager at least 3 business days in advance. "
        "Employees working remotely must be available during core hours of <b>10:00 AM to 3:00 PM (EST)</b> "
        "and maintain a stable internet connection with a minimum speed of 25 Mbps.", body))
    elements.append(Paragraph(
        "Eligible employees include all full-time staff who have completed their 90-day probation period. "
        "Contract workers and interns are not eligible for remote work arrangements. "
        "Remote work equipment (laptop, monitor, headset) is provided by the company. "
        "A home office stipend of <b>$75 per month</b> is available for internet and utility costs.", body))
    elements.append(Spacer(1, 10))

    # Section 2: Leave Policy
    elements.append(Paragraph("2. Leave Policy", h2))
    elements.append(Paragraph(
        "All full-time employees receive the following paid leave entitlements per calendar year:", body))

    leave_data = [
        ["Leave Type", "Days Per Year", "Carry Over"],
        ["Annual Leave", "20 days", "Up to 5 days"],
        ["Sick Leave", "12 days", "No carry over"],
        ["Personal Leave", "5 days", "No carry over"],
        ["Maternity Leave", "90 days", "N/A"],
        ["Paternity Leave", "30 days", "N/A"],
        ["Bereavement Leave", "5 days", "N/A"],
    ]
    table = Table(leave_data, colWidths=[2*inch, 1.5*inch, 1.5*inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        "Annual leave requests must be submitted at least <b>2 weeks in advance</b> through the HR portal. "
        "Sick leave exceeding 3 consecutive days requires a doctor's note. "
        "Unused annual leave (up to 5 days) carries over to the next year; remaining days are forfeited on December 31st.", body))
    elements.append(Spacer(1, 10))

    # Section 3: Expense Reimbursement
    elements.append(Paragraph("3. Expense Reimbursement Policy", h2))
    elements.append(Paragraph(
        "Employees may submit expense claims for pre-approved business expenses. "
        "All single expenses exceeding <b>$500</b> require prior written approval from the department head. "
        "The maximum monthly expense claim is <b>$2,000</b> for individual contributors and <b>$5,000</b> for managers.", body))
    elements.append(Paragraph(
        "Reimbursable categories include: travel (economy class flights, standard hotel), "
        "meals (up to $50/day for domestic, $75/day for international), client entertainment (up to $100/event), "
        "professional development courses (up to $1,500/year), and office supplies.", body))
    elements.append(Paragraph(
        "Expense reports must be submitted within <b>30 days</b> of the expense date. "
        "Late submissions beyond 60 days will not be reimbursed. "
        "All claims require itemized receipts. Processing time is 5-10 business days after approval.", body))
    elements.append(Spacer(1, 10))

    # Section 4: Code of Conduct
    elements.append(Paragraph("4. Code of Conduct", h2))
    elements.append(Paragraph(
        "All employees must adhere to the company's Code of Conduct at all times. "
        "Key principles include: integrity in all business dealings, respect for colleagues and clients, "
        "protection of confidential information, and compliance with all applicable laws and regulations.", body))
    elements.append(Paragraph(
        "Violations are categorized into three severity levels: "
        "<b>Level 1 (Minor)</b> — verbal warning, documented in file. Examples: occasional tardiness, dress code violations. "
        "<b>Level 2 (Moderate)</b> — written warning, possible suspension. Examples: unauthorized absence, misuse of company resources. "
        "<b>Level 3 (Severe)</b> — immediate termination. Examples: theft, harassment, fraud, disclosure of trade secrets.", body))
    elements.append(Spacer(1, 10))

    # Section 5: Performance Reviews
    elements.append(Paragraph("5. Performance Review Process", h2))
    elements.append(Paragraph(
        "Performance reviews are conducted <b>quarterly</b> in March, June, September, and December. "
        "Each review cycle consists of three phases: self-assessment (Week 1), "
        "peer feedback collection (Week 2), and manager evaluation meeting (Week 3).", body))
    elements.append(Paragraph(
        "The rating scale uses 5 levels: <b>1 = Needs Improvement</b>, <b>2 = Below Expectations</b>, "
        "<b>3 = Meets Expectations</b>, <b>4 = Exceeds Expectations</b>, <b>5 = Outstanding</b>. "
        "Employees rated 4 or 5 for two consecutive quarters are eligible for accelerated promotion consideration. "
        "Employees rated 1 for two consecutive quarters enter a 90-day Performance Improvement Plan (PIP).", body))
    elements.append(Paragraph(
        "Annual compensation reviews are tied to the December quarterly review. "
        "Merit increases range from <b>2% to 8%</b> based on overall rating. "
        "Promotion-based increases range from <b>10% to 20%</b> depending on the new role level.", body))

    doc.build(elements)
    print(f"  Created: {path} ({path.stat().st_size:,} bytes)")


def generate_financial_report_docx():
    """Generate a Q4 2025 financial report DOCX with tables."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    path = FIXTURES_DIR / "financial-report-q4-2025.docx"
    doc = Document()

    # Title
    title = doc.add_heading("Q4 2025 Financial Report", level=0)
    doc.add_paragraph("Acme Corporation | Fiscal Year Ending December 31, 2025").italic = True
    doc.add_paragraph("Prepared by: Finance Department | Date: January 15, 2026")
    doc.add_paragraph("")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Acme Corporation achieved strong financial results in Q4 2025, "
        "with total revenue of $2,400,000 representing an 18% year-over-year increase from Q4 2024 ($2,034,000). "
        "Net income for the quarter was $600,000, yielding a net margin of 25%. "
        "The company ended the year with $8.2 million in annual revenue, exceeding the board target of $7.8 million."
    )

    # Revenue Breakdown
    doc.add_heading("Revenue Breakdown", level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Revenue Stream", "Q4 2025", "Q4 2024", "YoY Change"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["SaaS Subscriptions", "$1,200,000", "$980,000", "+22%"],
        ["Professional Services", "$600,000", "$520,000", "+15%"],
        ["Enterprise Licenses", "$400,000", "$350,000", "+14%"],
        ["Support & Maintenance", "$150,000", "$134,000", "+12%"],
        ["Training & Certification", "$50,000", "$50,000", "0%"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "SaaS subscriptions remain the largest revenue driver at 50% of total Q4 revenue. "
        "The enterprise segment showed the strongest growth, driven by 3 new enterprise contracts "
        "signed in November 2025: Globex Inc ($150K/yr), Initech ($120K/yr), and Umbrella Corp ($130K/yr)."
    )

    # Expenses
    doc.add_heading("Operating Expenses", level=1)
    table2 = doc.add_table(rows=8, cols=3)
    table2.style = "Light Grid Accent 1"
    exp_headers = ["Category", "Q4 2025", "% of Revenue"]
    for i, h in enumerate(exp_headers):
        table2.rows[0].cells[i].text = h

    exp_data = [
        ["Salaries & Benefits", "$960,000", "40%"],
        ["Cloud Infrastructure", "$240,000", "10%"],
        ["Sales & Marketing", "$300,000", "12.5%"],
        ["Research & Development", "$180,000", "7.5%"],
        ["General & Administrative", "$72,000", "3%"],
        ["Depreciation", "$36,000", "1.5%"],
        ["Total Expenses", "$1,788,000", "74.5%"],
    ]
    for row_idx, row_data in enumerate(exp_data):
        for col_idx, val in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = val

    doc.add_paragraph("")

    # Net Income
    doc.add_heading("Net Income & Margins", level=1)
    doc.add_paragraph(
        "Q4 2025 net income was $612,000 (25.5% net margin) compared to Q4 2024 net income of $490,000 (24.1% net margin). "
        "The improvement in margin was primarily driven by operational efficiencies in cloud infrastructure, "
        "where costs grew only 12% despite 22% SaaS revenue growth."
    )

    # Cash Position
    doc.add_heading("Cash Position & Outlook", level=1)
    doc.add_paragraph(
        "Cash and cash equivalents as of December 31, 2025: $3,800,000. "
        "Accounts receivable: $720,000 (average collection period: 27 days). "
        "The company has no outstanding debt and an unused credit line of $2,000,000."
    )
    doc.add_paragraph(
        "Q1 2026 Guidance: Projected revenue of $2,100,000 to $2,300,000. "
        "Key initiatives include expansion into the APAC market (estimated $400K investment) "
        "and launch of the Acme AI Assistant product line (estimated $250K R&D spend). "
        "Expected headcount growth: 15 new hires across engineering, sales, and customer success."
    )

    doc.save(str(path))
    print(f"  Created: {path} ({path.stat().st_size:,} bytes)")


def generate_product_catalog_xlsx():
    """Generate a product catalog XLSX with multiple sheets."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    path = FIXTURES_DIR / "product-catalog.xlsx"
    wb = Workbook()

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")

    # Sheet 1: Products
    ws1 = wb.active
    ws1.title = "Products"
    products_headers = ["Product ID", "Product Name", "Category", "Description", "Weight (kg)", "Warranty (months)"]
    products_data = [
        ["WDG-001", "Widget Pro", "Widgets", "Industrial-grade widget with titanium casing, rated for 10,000+ cycles", 2.5, 24],
        ["WDG-002", "Widget Lite", "Widgets", "Lightweight consumer widget, ideal for home use, ergonomic grip design", 0.8, 12],
        ["WDG-003", "Widget Max", "Widgets", "Heavy-duty industrial widget, supports up to 500kg load capacity", 5.2, 36],
        ["GDG-001", "Gadget Alpha", "Gadgets", "Multi-function digital gadget with Bluetooth 5.0 and WiFi 6E", 0.3, 18],
        ["GDG-002", "Gadget Beta", "Gadgets", "Entry-level gadget with basic connectivity, USB-C charging", 0.2, 12],
        ["GDG-003", "Gadget Gamma", "Gadgets", "Premium gadget with OLED display, voice assistant, and GPS tracking", 0.45, 24],
        ["TLK-001", "ToolKit Basic", "Tools", "20-piece standard toolkit for home repairs, carrying case included", 3.1, 12],
        ["TLK-002", "ToolKit Professional", "Tools", "50-piece professional toolkit, chrome vanadium steel, lifetime warranty", 6.8, 60],
        ["ACC-001", "Power Adapter 100W", "Accessories", "Universal 100W USB-C power adapter, compatible with all Acme products", 0.15, 12],
        ["ACC-002", "Carrying Case XL", "Accessories", "Hard-shell carrying case, waterproof, fits all Widget and Gadget models", 1.2, 24],
    ]

    for col, h in enumerate(products_headers, 1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, row_data in enumerate(products_data, 2):
        for col_idx, val in enumerate(row_data, 1):
            ws1.cell(row=row_idx, column=col_idx, value=val)

    ws1.column_dimensions["A"].width = 12
    ws1.column_dimensions["B"].width = 22
    ws1.column_dimensions["C"].width = 14
    ws1.column_dimensions["D"].width = 60
    ws1.column_dimensions["E"].width = 14
    ws1.column_dimensions["F"].width = 18

    # Sheet 2: Pricing
    ws2 = wb.create_sheet("Pricing")
    pricing_headers = ["Product ID", "Product Name", "Tier 1 (1-9 units)", "Tier 2 (10-49 units)", "Tier 3 (50+ units)", "Currency"]
    pricing_data = [
        ["WDG-001", "Widget Pro", 49.99, 44.99, 39.99, "USD"],
        ["WDG-002", "Widget Lite", 24.99, 22.49, 19.99, "USD"],
        ["WDG-003", "Widget Max", 89.99, 80.99, 71.99, "USD"],
        ["GDG-001", "Gadget Alpha", 129.99, 116.99, 103.99, "USD"],
        ["GDG-002", "Gadget Beta", 59.99, 53.99, 47.99, "USD"],
        ["GDG-003", "Gadget Gamma", 199.99, 179.99, 159.99, "USD"],
        ["TLK-001", "ToolKit Basic", 34.99, 31.49, 27.99, "USD"],
        ["TLK-002", "ToolKit Professional", 79.99, 71.99, 63.99, "USD"],
        ["ACC-001", "Power Adapter 100W", 29.99, 26.99, 23.99, "USD"],
        ["ACC-002", "Carrying Case XL", 19.99, 17.99, 15.99, "USD"],
    ]

    for col, h in enumerate(pricing_headers, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, row_data in enumerate(pricing_data, 2):
        for col_idx, val in enumerate(row_data, 1):
            ws2.cell(row=row_idx, column=col_idx, value=val)

    ws2.column_dimensions["A"].width = 12
    ws2.column_dimensions["B"].width = 22
    ws2.column_dimensions["C"].width = 20
    ws2.column_dimensions["D"].width = 22
    ws2.column_dimensions["E"].width = 20
    ws2.column_dimensions["F"].width = 12

    # Sheet 3: Inventory
    ws3 = wb.create_sheet("Inventory")
    inv_headers = ["Product ID", "Product Name", "Warehouse A (NY)", "Warehouse B (LA)", "Warehouse C (CHI)", "Total Stock", "Reorder Level", "Status"]
    inv_data = [
        ["WDG-001", "Widget Pro", 150, 80, 70, 300, 50, "In Stock"],
        ["WDG-002", "Widget Lite", 200, 120, 100, 420, 75, "In Stock"],
        ["WDG-003", "Widget Max", 25, 15, 10, 50, 20, "Low Stock"],
        ["GDG-001", "Gadget Alpha", 85, 60, 45, 190, 30, "In Stock"],
        ["GDG-002", "Gadget Beta", 300, 200, 150, 650, 100, "In Stock"],
        ["GDG-003", "Gadget Gamma", 10, 5, 5, 20, 15, "Critical"],
        ["TLK-001", "ToolKit Basic", 75, 50, 40, 165, 25, "In Stock"],
        ["TLK-002", "ToolKit Professional", 40, 30, 20, 90, 20, "In Stock"],
        ["ACC-001", "Power Adapter 100W", 500, 300, 250, 1050, 100, "In Stock"],
        ["ACC-002", "Carrying Case XL", 0, 0, 0, 0, 25, "Out of Stock"],
    ]

    for col, h in enumerate(inv_headers, 1):
        cell = ws3.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, row_data in enumerate(inv_data, 2):
        for col_idx, val in enumerate(row_data, 1):
            ws3.cell(row=row_idx, column=col_idx, value=val)

    ws3.column_dimensions["A"].width = 12
    ws3.column_dimensions["B"].width = 22
    ws3.column_dimensions["C"].width = 18
    ws3.column_dimensions["D"].width = 18
    ws3.column_dimensions["E"].width = 18
    ws3.column_dimensions["F"].width = 14
    ws3.column_dimensions["G"].width = 16
    ws3.column_dimensions["H"].width = 14

    wb.save(str(path))
    print(f"  Created: {path} ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    print("Generating sample test files for RAG pipeline testing...\n")
    generate_company_policy_pdf()
    generate_financial_report_docx()
    generate_product_catalog_xlsx()
    print(f"\nAll files saved to: {FIXTURES_DIR}")
    print("Use these files with: curl -X POST -H 'Authorization: Bearer $TOKEN' -F 'file=@<path>' $BASE/documents/upload")
